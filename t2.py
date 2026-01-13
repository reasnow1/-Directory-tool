import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document
from docx.shared import Inches
import threading
from datetime import datetime

class DirectoryExtractor:
    def __init__(self, root):
        self.root = root
        self.root.title("目录文件名提取工具")
        self.root.geometry("830x600")
        
        # 设置字体
        self.font = ("Microsoft YaHei", 10)
        
        # 创建界面元素
        self.create_widgets()
        
    def create_widgets(self):
        # 目录选择区域
        dir_frame = tk.Frame(self.root)
        dir_frame.pack(pady=5, padx=20, fill="x")
        
        tk.Label(dir_frame, text="目录路径:", font=self.font).pack(side="left")
        
        # 目录输入框
        self.dir_var = tk.StringVar()
        self.dir_entry = tk.Entry(dir_frame, textvariable=self.dir_var, 
                                  font=self.font, width=50)
        self.dir_entry.pack(side="left", padx=10)
        
        # 浏览按钮
        browse_btn = tk.Button(dir_frame, text="浏览", font=self.font,
                              command=self.browse_directory, width=8)
        browse_btn.pack(side="left", padx=10)
        
        
        self.read_btn = tk.Button(dir_frame, text="读取目录", font=self.font,
                                 command=self.read_directory, width=8)
        self.read_btn.pack(side="left", padx=10)
        
        self.export_btn = tk.Button(dir_frame, text="导出为Word", font=self.font,
                                   command=self.export_to_word, width=10)
        self.export_btn.pack(side="left", padx=10)
        
        # 选项区域
        options_frame = tk.Frame(self.root)
        options_frame.pack(pady=1, padx=20, anchor='w')
        
        # 扩展名选项
        self.ext_var = tk.BooleanVar(value=True)
        self.ext_check = tk.Checkbutton(options_frame, text="显示扩展名", 
                                       variable=self.ext_var, font=self.font)
        self.ext_check.pack(side="left", padx=20)
        
        # 子目录选项
        self.subdir_var = tk.BooleanVar(value=True)
        self.subdir_check = tk.Checkbutton(options_frame, text="包括子目录", 
                                          variable=self.subdir_var, font=self.font)
        self.subdir_check.pack(side="left", padx=20)
        
        # 文件类型筛选选项
        self.filter_var = tk.StringVar(value="所有文件")

        
        tk.Label(options_frame, text="文件类型:", font=self.font).pack(side="left")
        
        filter_options = ["所有文件", "图片文件", "文档文件", "视频文件", "音频文件", "压缩文件"]
        filter_menu = tk.OptionMenu(options_frame, self.filter_var, *filter_options)
        filter_menu.config(font=self.font)
        filter_menu.pack(side="left", padx=10)
        
        # 输出文本框和滚动条
        output_frame = tk.Frame(self.root)
        output_frame.pack(pady=1, padx=20, fill="both", expand=True)
        
        tk.Label(output_frame, text="文件列表:", font=self.font).pack(anchor="w")
        
        # 文本框和滚动条
        text_scroll = tk.Scrollbar(output_frame)
        text_scroll.pack(side="right", fill="y")
        
        self.output_text = tk.Text(output_frame, font=("Consolas", 9), 
                                   yscrollcommand=text_scroll.set,
                                   wrap="none", height=15)
        self.output_text.pack(fill="both", expand=True)
        
        text_scroll.config(command=self.output_text.yview)
        
        # 状态栏
        self.status_var = tk.StringVar(value="就绪")
        status_bar = tk.Label(self.root, textvariable=self.status_var, 
                             font=self.font, bd=1, relief=tk.SUNKEN, anchor="w")
        status_bar.pack(side="bottom", fill="x")
        
    def browse_directory(self):
        """浏览目录"""
        directory = filedialog.askdirectory(title="选择目录")
        if directory:
            self.dir_var.set(directory)
            
    def read_directory(self):
        """读取目录"""
        directory = self.dir_var.get()
        
        if not directory or not os.path.exists(directory):
            messagebox.showerror("错误", "请选择有效的目录路径！")
            return
            
        # 禁用按钮，避免重复操作
        self.read_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.status_var.set("正在读取目录...")
        
        # 在新线程中执行读取操作
        thread = threading.Thread(target=self._read_directory_thread, args=(directory,))
        thread.daemon = True
        thread.start()
        
    def _read_directory_thread(self, directory):
        """读取目录的线程函数"""
        try:
            show_ext = self.ext_var.get()
            include_subdir = self.subdir_var.get()
            file_filter = self.filter_var.get()
            
            # 清空输出框
            self.output_text.delete(1.0, tk.END)
            
            # 获取文件列表
            file_list = []
            
            if include_subdir:
                for root_dir, dirs, files in os.walk(directory):
                    for file in files:
                        file_path = os.path.join(root_dir, file)
                        relative_path = os.path.relpath(file_path, directory)
                        file_list.append((relative_path, file_path))
            else:
                for item in os.listdir(directory):
                    item_path = os.path.join(directory, item)
                    if os.path.isfile(item_path):
                        file_list.append((item, item_path))
            
            # 应用文件类型筛选
            file_list = self.filter_files(file_list, file_filter)
            
            # 显示结果
            self.display_file_list(file_list, directory, show_ext)
            
            # 更新状态
            self.root.after(0, lambda: self.status_var.set(f"找到 {len(file_list)} 个文件"))
            
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("错误", f"读取目录时出错: {str(e)}"))
            self.root.after(0, lambda: self.status_var.set("读取失败"))
        finally:
            # 重新启用按钮
            self.root.after(0, lambda: self.read_btn.config(state="normal"))
            self.root.after(0, lambda: self.export_btn.config(state="normal"))
            
    def filter_files(self, file_list, file_filter):
        """根据选择的文件类型筛选文件"""
        if file_filter == "所有文件":
            return file_list
            
        extensions = {
            "图片文件": ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            "文档文件": ['.txt', '.doc', '.docx', '.pdf', '.xls', '.xlsx', '.ppt', '.pptx'],
            "视频文件": ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.mpg'],
            "音频文件": ['.mp3', '.wav', '.flac', '.aac', '.wma', '.ogg'],
            "压缩文件": ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2']
        }
        
        filter_exts = extensions.get(file_filter, [])
        filtered_list = []
        
        for relative_path, full_path in file_list:
            ext = os.path.splitext(full_path)[1].lower()
            if ext in filter_exts:
                filtered_list.append((relative_path, full_path))
                
        return filtered_list
        
    def display_file_list(self, file_list, base_dir, show_ext):
        """在文本框中显示文件列表"""
        if not file_list:
            self.output_text.insert(tk.END, "未找到文件")
            return
            
        # 添加标题
        self.output_text.insert(tk.END, f"目录: {base_dir}\n")
        self.output_text.insert(tk.END, f"扫描时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        self.output_text.insert(tk.END, "="*60 + "\n\n")
        
        # 添加文件列表
        for i, (relative_path, full_path) in enumerate(file_list, 1):
            if not show_ext:
                # 移除扩展名
                relative_path = os.path.splitext(relative_path)[0]
                
            file_size = os.path.getsize(full_path)
            size_str = self.format_file_size(file_size)
            
            # 显示文件信息
            self.output_text.insert(tk.END, f"{i:4d}. {relative_path}")
            self.output_text.insert(tk.END, f" ({size_str})\n")
            
        self.output_text.see(1.0)  # 滚动到顶部
        
    def format_file_size(self, size_bytes):
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} TB"
        
    def export_to_word(self):
        """导出为Word文档"""
        if not self.output_text.get(1.0, tk.END).strip():
            messagebox.showwarning("警告", "请先读取目录内容！")
            return
            
        # 选择保存位置
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            title="保存Word文档"
        )
        
        if not file_path:
            return
            
        # 禁用按钮，避免重复操作
        self.read_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.status_var.set("正在导出Word文档...")
        
        # 在新线程中执行导出操作
        thread = threading.Thread(target=self._export_to_word_thread, args=(file_path,))
        thread.daemon = True
        thread.start()
        
    def _export_to_word_thread(self, file_path):
        """导出Word文档的线程函数"""
        try:
            # 获取文本框内容
            content = self.output_text.get(1.0, tk.END)
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading('目录文件列表', 0)
            
            # 添加时间信息
            current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            doc.add_paragraph(f'生成时间: {current_time}')
            
            # 添加目录信息
            lines = content.split('\n')
            for line in lines:
                if line.strip():  # 跳过空行
                    if '=' in line:  # 分隔线
                        doc.add_paragraph(line)
                    elif line.startswith('目录:') or line.startswith('扫描时间:'):
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph(line)
            
            # 保存文档
            doc.save(file_path)
            
            self.root.after(0, lambda: messagebox.showinfo("成功", f"已导出到: {file_path}"))
            self.root.after(0, lambda: self.status_var.set("导出完成"))
            
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("错误", f"导出失败: {str(e)}"))
            self.root.after(0, lambda: self.status_var.set("导出失败"))
        finally:
            # 重新启用按钮
            self.root.after(0, lambda: self.read_btn.config(state="normal"))
            self.root.after(0, lambda: self.export_btn.config(state="normal"))

def main():
    root = tk.Tk()
    app = DirectoryExtractor(root)
    
    # 设置窗口图标（可选）
    try:
        root.iconbitmap(default='icon.ico')  # 如果有图标文件的话
    except:
        pass
    
    root.mainloop()

if __name__ == "__main__":
    main()